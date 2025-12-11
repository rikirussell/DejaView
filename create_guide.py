from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Title
title = doc.add_heading('DejaView App Store Publishing Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Complete Step-by-Step Guide for Apple App Store & Google Play Store')
doc.add_paragraph('')

# Table of Contents
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph('1. Prerequisites & Requirements')
doc.add_paragraph('2. Building Your App with EAS (Expo Application Services)')
doc.add_paragraph('3. Apple App Store Submission')
doc.add_paragraph('4. Google Play Store Submission')
doc.add_paragraph('5. App Store Assets Checklist')
doc.add_paragraph('6. Common Issues & Troubleshooting')
doc.add_paragraph('')

# Section 1
doc.add_heading('1. Prerequisites & Requirements', level=1)

doc.add_heading('For Both Platforms:', level=2)
doc.add_paragraph('• Expo account (free at expo.dev)')
doc.add_paragraph('• EAS CLI installed: npm install -g eas-cli')
doc.add_paragraph('• App icons (1024x1024 PNG, no transparency for iOS)')
doc.add_paragraph('• Screenshots for various device sizes')
doc.add_paragraph('• Privacy Policy URL (required by both stores)')
doc.add_paragraph('• App description (short and long versions)')

doc.add_heading('For Apple App Store:', level=2)
doc.add_paragraph('• Apple Developer Account ($99/year) - developer.apple.com')
doc.add_paragraph('• Mac computer (required for final upload via Transporter or Xcode)')
doc.add_paragraph('• App Store Connect access')

doc.add_heading('For Google Play Store:', level=2)
doc.add_paragraph('• Google Play Developer Account ($25 one-time) - play.google.com/console')
doc.add_paragraph('• Google Play Console access')
doc.add_paragraph('')

# Section 2
doc.add_heading('2. Building Your App with EAS', level=1)

doc.add_heading('Step 2.1: Install EAS CLI', level=2)
doc.add_paragraph('Open your terminal and run:')
code = doc.add_paragraph('npm install -g eas-cli')
code.style = 'Quote'

doc.add_heading('Step 2.2: Login to Expo', level=2)
code = doc.add_paragraph('eas login')
code.style = 'Quote'
doc.add_paragraph('Enter your Expo account credentials.')

doc.add_heading('Step 2.3: Configure EAS Build', level=2)
doc.add_paragraph('Navigate to your project directory and run:')
code = doc.add_paragraph('cd /path/to/your/frontend\neas build:configure')
code.style = 'Quote'
doc.add_paragraph('This creates an eas.json file with build configurations.')

doc.add_heading('Step 2.4: Update app.json', level=2)
doc.add_paragraph('Ensure your app.json has these required fields:')
config_text = '''{
  "expo": {
    "name": "DejaView",
    "slug": "dejaview",
    "version": "1.0.0",
    "orientation": "default",
    "icon": "./assets/images/icon.png",
    "splash": {
      "image": "./assets/images/splash.png",
      "resizeMode": "contain",
      "backgroundColor": "#0c0c0c"
    },
    "ios": {
      "supportsTablet": true,
      "bundleIdentifier": "com.yourcompany.dejaview",
      "buildNumber": "1",
      "infoPlist": {
        "NSCameraUsageDescription": "DejaView needs camera access to capture photos and videos with overlay alignment.",
        "NSPhotoLibraryUsageDescription": "DejaView needs photo library access to load overlay images and save captured media.",
        "NSMicrophoneUsageDescription": "DejaView needs microphone access to record video with audio."
      }
    },
    "android": {
      "adaptiveIcon": {
        "foregroundImage": "./assets/images/adaptive-icon.png",
        "backgroundColor": "#0c0c0c"
      },
      "package": "com.yourcompany.dejaview",
      "versionCode": 1,
      "permissions": [
        "CAMERA",
        "READ_EXTERNAL_STORAGE",
        "WRITE_EXTERNAL_STORAGE",
        "RECORD_AUDIO"
      ]
    }
  }
}'''
code = doc.add_paragraph(config_text)
code.style = 'Quote'

doc.add_heading('Step 2.5: Build for iOS', level=2)
code = doc.add_paragraph('eas build --platform ios --profile production')
code.style = 'Quote'
doc.add_paragraph('• First time: EAS will ask to create iOS credentials (let EAS manage them)')
doc.add_paragraph('• Build takes 15-30 minutes')
doc.add_paragraph('• Download the .ipa file when complete')

doc.add_heading('Step 2.6: Build for Android', level=2)
code = doc.add_paragraph('eas build --platform android --profile production')
code.style = 'Quote'
doc.add_paragraph('• First time: EAS will create a keystore (let EAS manage it)')
doc.add_paragraph('• Build takes 10-20 minutes')
doc.add_paragraph('• Download the .aab file when complete')
doc.add_paragraph('')

# Section 3
doc.add_heading('3. Apple App Store Submission', level=1)

doc.add_heading('Step 3.1: Create App Store Connect Listing', level=2)
doc.add_paragraph('1. Go to appstoreconnect.apple.com')
doc.add_paragraph('2. Click "My Apps" → "+" → "New App"')
doc.add_paragraph('3. Fill in:')
doc.add_paragraph('   • Platform: iOS')
doc.add_paragraph('   • Name: DejaView')
doc.add_paragraph('   • Primary Language: English (or your language)')
doc.add_paragraph('   • Bundle ID: Select the one matching your app.json')
doc.add_paragraph('   • SKU: dejaview-ios-001 (unique identifier)')
doc.add_paragraph('   • User Access: Full Access')

doc.add_heading('Step 3.2: Prepare App Information', level=2)
doc.add_paragraph('In App Store Connect, fill in:')
doc.add_paragraph('')
doc.add_paragraph('App Information Tab:')
doc.add_paragraph('• Subtitle: "Perfect photo alignment every time"')
doc.add_paragraph('• Category: Photo & Video')
doc.add_paragraph('• Content Rights: Confirm you own or have rights')
doc.add_paragraph('• Age Rating: Complete the questionnaire (likely 4+)')
doc.add_paragraph('')
doc.add_paragraph('Pricing and Availability:')
doc.add_paragraph('• Price: Free (or your chosen price)')
doc.add_paragraph('• Availability: All countries or select specific ones')

doc.add_heading('Step 3.3: Prepare App Store Listing', level=2)
doc.add_paragraph('Version Information:')
doc.add_paragraph('')
doc.add_paragraph('Screenshots Required:')
doc.add_paragraph('• 6.7" Display (iPhone 14 Pro Max): 1290 x 2796 pixels')
doc.add_paragraph('• 6.5" Display (iPhone 11 Pro Max): 1284 x 2778 pixels')
doc.add_paragraph('• 5.5" Display (iPhone 8 Plus): 1242 x 2208 pixels')
doc.add_paragraph('• iPad Pro 12.9" (if supporting tablet): 2048 x 2732 pixels')
doc.add_paragraph('')
doc.add_paragraph('Promotional Text (170 chars):')
doc.add_paragraph('"Recreate your favorite shots perfectly! Load any photo as a transparent overlay on your camera to align and capture matching compositions."')
doc.add_paragraph('')
doc.add_paragraph('Description:')
description_text = '''DejaView helps you perfectly recreate or align shots by loading a previously captured photo as a semi-transparent onion-skin overlay inside the live camera viewfinder.

FEATURES:
• Load any photo as a semi-transparent overlay
• Adjust overlay opacity with intuitive slider
• Pinch to zoom, drag to position, rotate to align
• Capture photos aligned with your reference
• Record video with overlay visible
• Save directly to your photo library
• Support for portrait and landscape orientations
• Grid guides for precise composition

PERFECT FOR:
• Before/after comparisons
• Time-lapse sequences
• Recreating classic photos
• Real estate photography
• Progress photos (fitness, construction, etc.)
• Creative double exposures
• Educational photography

Download DejaView and never miss the perfect alignment again!'''
doc.add_paragraph(description_text)
doc.add_paragraph('')
doc.add_paragraph('Keywords (100 chars max, comma-separated):')
doc.add_paragraph('"overlay,camera,alignment,photo,recreate,onion skin,composition,reference,before after,timelapse"')

doc.add_heading('Step 3.4: Upload Your Build', level=2)
doc.add_paragraph('Option A - Using EAS Submit (Recommended):')
code = doc.add_paragraph('eas submit --platform ios')
code.style = 'Quote'
doc.add_paragraph('• Follow prompts to select your build')
doc.add_paragraph('• EAS will upload directly to App Store Connect')
doc.add_paragraph('')
doc.add_paragraph('Option B - Using Transporter (Mac only):')
doc.add_paragraph('1. Download Transporter from Mac App Store')
doc.add_paragraph('2. Sign in with your Apple ID')
doc.add_paragraph('3. Drag and drop your .ipa file')
doc.add_paragraph('4. Click "Deliver"')

doc.add_heading('Step 3.5: Submit for Review', level=2)
doc.add_paragraph('1. In App Store Connect, go to your app')
doc.add_paragraph('2. Click "+ Version or Platform" if needed')
doc.add_paragraph('3. Select your uploaded build')
doc.add_paragraph('4. Complete App Review Information:')
doc.add_paragraph('   • Sign-In Information: N/A (no login required)')
doc.add_paragraph('   • Contact Information: Your details')
doc.add_paragraph('   • Notes: "This app uses the camera to capture photos/videos with overlay alignment. No account required."')
doc.add_paragraph('5. Click "Submit for Review"')
doc.add_paragraph('')
doc.add_paragraph('Review typically takes 24-48 hours.')
doc.add_paragraph('')

# Section 4
doc.add_heading('4. Google Play Store Submission', level=1)

doc.add_heading('Step 4.1: Create Google Play Console Listing', level=2)
doc.add_paragraph('1. Go to play.google.com/console')
doc.add_paragraph('2. Click "Create app"')
doc.add_paragraph('3. Fill in:')
doc.add_paragraph('   • App name: DejaView')
doc.add_paragraph('   • Default language: English')
doc.add_paragraph('   • App or game: App')
doc.add_paragraph('   • Free or paid: Free (or Paid)')
doc.add_paragraph('4. Accept declarations and click "Create app"')

doc.add_heading('Step 4.2: Set Up Your Store Listing', level=2)
doc.add_paragraph('Go to "Grow" → "Store presence" → "Main store listing"')
doc.add_paragraph('')
doc.add_paragraph('App Details:')
doc.add_paragraph('• Short description (80 chars): "Load photo overlays on your camera to perfectly recreate and align shots."')
doc.add_paragraph('• Full description: (Same as iOS description above)')
doc.add_paragraph('')
doc.add_paragraph('Graphics:')
doc.add_paragraph('• App icon: 512 x 512 PNG')
doc.add_paragraph('• Feature graphic: 1024 x 500 PNG (promotional banner)')
doc.add_paragraph('• Screenshots: At least 2, recommended 8')
doc.add_paragraph('  - Phone: 16:9 or 9:16 aspect ratio, min 320px, max 3840px')
doc.add_paragraph('  - Tablet (if supporting): 16:9 or 9:16 aspect ratio')

doc.add_heading('Step 4.3: Complete Store Settings', level=2)
doc.add_paragraph('App Content (Required):')
doc.add_paragraph('')
doc.add_paragraph('1. Privacy Policy:')
doc.add_paragraph('   • Add your privacy policy URL')
doc.add_paragraph('   • Must be publicly accessible')
doc.add_paragraph('')
doc.add_paragraph('2. App Access:')
doc.add_paragraph('   • Select "All functionality is available without special access"')
doc.add_paragraph('')
doc.add_paragraph('3. Ads:')
doc.add_paragraph('   • Select "No, my app does not contain ads"')
doc.add_paragraph('')
doc.add_paragraph('4. Content Rating:')
doc.add_paragraph('   • Complete the questionnaire')
doc.add_paragraph('   • DejaView should receive "Everyone" rating')
doc.add_paragraph('')
doc.add_paragraph('5. Target Audience:')
doc.add_paragraph('   • Select age groups (likely 13+ or All ages)')
doc.add_paragraph('')
doc.add_paragraph('6. News Apps:')
doc.add_paragraph('   • Select "No"')
doc.add_paragraph('')
doc.add_paragraph('7. Data Safety:')
doc.add_paragraph('   • Complete the form about data collection')
doc.add_paragraph('   • DejaView collects: Photos/Videos (stored locally only)')
doc.add_paragraph('   • Data is not shared with third parties')

doc.add_heading('Step 4.4: Upload Your Build', level=2)
doc.add_paragraph('Option A - Using EAS Submit (Recommended):')
code = doc.add_paragraph('eas submit --platform android')
code.style = 'Quote'
doc.add_paragraph('• First time: Create a Google Service Account key')
doc.add_paragraph('• Follow EAS prompts to complete upload')
doc.add_paragraph('')
doc.add_paragraph('Option B - Manual Upload:')
doc.add_paragraph('1. Go to "Release" → "Production"')
doc.add_paragraph('2. Click "Create new release"')
doc.add_paragraph('3. Upload your .aab file')
doc.add_paragraph('4. Add release notes: "Initial release of DejaView"')
doc.add_paragraph('5. Click "Review release"')

doc.add_heading('Step 4.5: Submit for Review', level=2)
doc.add_paragraph('1. Ensure all sections show green checkmarks')
doc.add_paragraph('2. Go to "Release" → "Production"')
doc.add_paragraph('3. Click "Start rollout to Production"')
doc.add_paragraph('4. Confirm the rollout')
doc.add_paragraph('')
doc.add_paragraph('Review typically takes a few hours to 7 days.')
doc.add_paragraph('')

# Section 5
doc.add_heading('5. App Store Assets Checklist', level=1)

doc.add_heading('Icons', level=2)
doc.add_paragraph('□ iOS App Icon: 1024 x 1024 PNG (no alpha/transparency)')
doc.add_paragraph('□ Android App Icon: 512 x 512 PNG')
doc.add_paragraph('□ Android Adaptive Icon: 432 x 432 PNG with safe zone')

doc.add_heading('Screenshots', level=2)
doc.add_paragraph('iOS (minimum 3 per device size):')
doc.add_paragraph('□ 6.7" (1290 x 2796)')
doc.add_paragraph('□ 6.5" (1284 x 2778)')
doc.add_paragraph('□ 5.5" (1242 x 2208)')
doc.add_paragraph('□ iPad 12.9" (2048 x 2732) - if supporting tablets')
doc.add_paragraph('')
doc.add_paragraph('Android (minimum 2):')
doc.add_paragraph('□ Phone screenshots (16:9 or 9:16)')
doc.add_paragraph('□ Feature graphic (1024 x 500)')

doc.add_heading('Text Content', level=2)
doc.add_paragraph('□ App name (30 chars)')
doc.add_paragraph('□ Subtitle/Short description (30-80 chars)')
doc.add_paragraph('□ Full description (4000 chars max)')
doc.add_paragraph('□ Keywords (iOS: 100 chars)')
doc.add_paragraph('□ What\'s New / Release notes')
doc.add_paragraph('□ Privacy Policy URL')
doc.add_paragraph('□ Support URL')
doc.add_paragraph('')

# Section 6
doc.add_heading('6. Common Issues & Troubleshooting', level=1)

doc.add_heading('iOS Issues', level=2)
doc.add_paragraph('Issue: "Missing Compliance" warning')
doc.add_paragraph('Solution: Add to app.json under ios:')
code = doc.add_paragraph('"infoPlist": {\n  "ITSAppUsesNonExemptEncryption": false\n}')
code.style = 'Quote'
doc.add_paragraph('')
doc.add_paragraph('Issue: Build rejected for missing permissions description')
doc.add_paragraph('Solution: Ensure all NSxxxUsageDescription keys are in your app.json')
doc.add_paragraph('')
doc.add_paragraph('Issue: Screenshots rejected')
doc.add_paragraph('Solution: Ensure screenshots show actual app functionality, no device frames with status bars showing incorrect time')

doc.add_heading('Android Issues', level=2)
doc.add_paragraph('Issue: "App not compatible with any devices"')
doc.add_paragraph('Solution: Check your android permissions in app.json and eas.json build settings')
doc.add_paragraph('')
doc.add_paragraph('Issue: "Deceptive behavior" policy violation')
doc.add_paragraph('Solution: Ensure all camera/photo permissions are properly declared and used as described')

doc.add_heading('General Tips', level=2)
doc.add_paragraph('• Test your app thoroughly before submission')
doc.add_paragraph('• Use TestFlight (iOS) and Internal Testing (Android) first')
doc.add_paragraph('• Respond quickly to any review feedback')
doc.add_paragraph('• Keep your privacy policy up to date')
doc.add_paragraph('• Monitor crash reports and user feedback after launch')
doc.add_paragraph('')

# Final section
doc.add_heading('Quick Reference Commands', level=1)
commands = '''# Login to EAS
eas login

# Configure EAS for your project
eas build:configure

# Build for iOS
eas build --platform ios --profile production

# Build for Android  
eas build --platform android --profile production

# Submit to iOS App Store
eas submit --platform ios

# Submit to Google Play Store
eas submit --platform android

# Build for both platforms
eas build --platform all --profile production'''
code = doc.add_paragraph(commands)
code.style = 'Quote'

doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('Document created for DejaView App')
doc.add_paragraph('Good luck with your app store submissions!')

# Save document
doc.save('/app/frontend/DejaView_App_Store_Publishing_Guide.docx')
print("Document created successfully!")
